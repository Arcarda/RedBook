from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_docx_import():
    input_file = 'Red_Book_Manu_FINAL.md'
    output_file = 'Red_Book_Manu_Import_Perfect.docx'
    
    doc = Document()
    
    # -----------------------------------------------------
    # Setup Styles Matches Affinity Mapping
    # -----------------------------------------------------
    # We use built-in styles where possible as per the guide:
    # Normal -> Body Text
    # Heading 1 -> Part Title
    # Heading 2 -> Section Header
    # Heading 3 -> Subsection Header
    # Quote -> Blockquote
    
    styles = doc.styles
    
    # Custom Style for Operator's Orders if needed, or we just use a specific one
    # Let's create a custom style "OperatorsOrder" to be safe and distinct
    try:
        style = styles.add_style('OperatorsOrder', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']
        font = style.font
        font.name = 'Arial' # Distinctive for the draft
        font.bold = True
    except:
        pass # Style might exist

    # -----------------------------------------------------
    # Markdown Parser
    # -----------------------------------------------------
    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    current_context = "Normal" # Normal, Code, OperatorOrder
    
    # Regex Patterns
    re_h1 = re.compile(r'^#\s+(.*)')
    re_h2 = re.compile(r'^##\s+(.*)')
    re_h3 = re.compile(r'^###\s+(.*)')
    re_blockquote = re.compile(r'^>\s+(.*)')
    re_bold = re.compile(r'\*\*(.*?)\*\*')
    re_italic = re.compile(r'\*(.*?)\*')
    
    operator_mode = False
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Empty line
            continue
            
        # Check for Operator's Orders header
        if "OPERATOR'S ORDERS" in line:
            operator_mode = True
            p = doc.add_paragraph(line, style='Heading 3') # Keep the header as H3
            continue
            
        # Check for leaving Operator's Orders (usually a horizontal rule or new header)
        if line.startswith('---') or line.startswith('#'):
            operator_mode = False
            
        # Headers
        m_h1 = re_h1.match(line)
        if m_h1:
            doc.add_heading(m_h1.group(1), level=1)
            continue
            
        m_h2 = re_h2.match(line)
        if m_h2:
            doc.add_heading(m_h2.group(1), level=2)
            continue
            
        m_h3 = re_h3.match(line)
        if m_h3:
            doc.add_heading(m_h3.group(1), level=3)
            continue
            
        # Blockquotes
        m_quote = re_blockquote.match(line)
        if m_quote:
            text = m_quote.group(1)
            p = doc.add_paragraph(text, style='Quote')
            # Apply formatting
            apply_formatting(p, text)
            continue
            
        # Normal Text / Operator List
        if operator_mode:
            # If it looks like a list item
            if line[0].isdigit() or line.startswith('-'):
                p = doc.add_paragraph(line, style='OperatorsOrder')
            else:
                p = doc.add_paragraph(line, style='OperatorsOrder')
        else:
            # Standard Paragraph
            p = doc.add_paragraph(style='Normal')
            apply_formatting(p, line)

    doc.save(output_file)
    print(f"DOCX created at: {output_file}")

def apply_formatting(paragraph, text):
    """
    Simple parser to handle **bold** and *italic* within a paragraph string.
    Splits by bold, then italic.
    """
    # This is a basic implementation. For nested or complex, full Markdown parser is needed.
    # Python-docx requires adding 'runs'.
    
    # Strategy: Split by **
    # "Hello **bold** world" -> ["Hello ", "bold", " world"]
    
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        bold = False
        if part.startswith('**') and part.endswith('**'):
            bold = True
            content = part[2:-2]
        else:
            content = part
            
        # Now handle italics inside this part
        subparts = re.split(r'(\*.*?\*)', content)
        for subpart in subparts:
            italic = False
            if subpart.startswith('*') and subpart.endswith('*') and len(subpart) > 2:
                italic = True
                clean_text = subpart[1:-1]
            else:
                clean_text = subpart
            
            if clean_text:
                run = paragraph.add_run(clean_text)
                if bold: run.bold = True
                if italic: run.italic = True

if __name__ == '__main__':
    create_docx_import()
