from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_journal_docx_import():
    input_file = 'Field_Journal_Manu.md'
    output_file = 'Field_Journal_Import_Perfect.docx'
    
    doc = Document()
    
    # -----------------------------------------------------
    # Setup Styles Matches Affinity Mapping (Consistent with Red Book)
    # -----------------------------------------------------
    # Normal -> Body Text
    # Heading 1 -> Part Title
    # Heading 2 -> Section Header
    # Heading 3 -> Subsection Header
    # Quote -> Blockquote
    
    styles = doc.styles
    
    # -----------------------------------------------------
    # Markdown Parser
    # -----------------------------------------------------
    try:
        with open(input_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except FileNotFoundError:
        print(f"Error: {input_file} not found.")
        return
        
    # Regex Patterns
    re_h1 = re.compile(r'^#\s+(.*)')
    re_h2 = re.compile(r'^##\s+(.*)')
    re_h3 = re.compile(r'^###\s+(.*)')
    re_blockquote = re.compile(r'^>\s+(.*)')
    
    # List state
    in_list = False
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Empty line
            continue
            
        # Headers
        m_h1 = re_h1.match(line)
        if m_h1:
            doc.add_heading(m_h1.group(1), level=1)
            continue
            
        m_h2 = re_h2.match(line)
        if m_h2:
            doc.add_heading(m_h2.group(1), level=2)
            continue
            
        m_h3 = re_h3.match(line)
        if m_h3:
            doc.add_heading(m_h3.group(1), level=3)
            continue
            
        # Blockquotes
        m_quote = re_blockquote.match(line)
        if m_quote:
            text = m_quote.group(1)
            p = doc.add_paragraph(text, style='Quote')
            apply_formatting(p, text)
            continue
            
        # Lists (Basic)
        if line.startswith('* ') or line.startswith('- '):
            clean_line = line[2:]
            p = doc.add_paragraph(clean_line, style='List Bullet')
            apply_formatting(p, clean_line)
            continue
            
        if line[0].isdigit() and line[1] == '.':
            # 1. Item
            clean_line = line.split('.', 1)[1].strip()
            p = doc.add_paragraph(clean_line, style='List Number')
            apply_formatting(p, clean_line)
            continue
            
        # Standard Paragraph
        p = doc.add_paragraph(style='Normal')
        apply_formatting(p, line)

    doc.save(output_file)
    print(f"DOCX created at: {output_file}")

def apply_formatting(paragraph, text):
    """
    Simple parser to handle **bold** and *italic* within a paragraph string.
    """
    # Strategy: Split by **
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        bold = False
        if part.startswith('**') and part.endswith('**'):
            bold = True
            content = part[2:-2]
        else:
            content = part
            
        # Now handle italics inside this part
        subparts = re.split(r'(\*.*?\*)', content)
        for subpart in subparts:
            italic = False
            if subpart.startswith('*') and subpart.endswith('*') and len(subpart) > 2:
                italic = True
                clean_text = subpart[1:-1]
            else:
                clean_text = subpart
            
            if clean_text:
                run = paragraph.add_run(clean_text)
                if bold: run.bold = True
                if italic: run.italic = True

if __name__ == '__main__':
    create_journal_docx_import()
